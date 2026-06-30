import json
import sys
from pathlib import Path

from docx import Document


def clean(text: str) -> str:
    return " ".join(text.replace("\u00ad", "").split())


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: extract_book.py INPUT.docx OUTPUT.json")

    source = Path(sys.argv[1])
    destination = Path(sys.argv[2])
    document = Document(source)

    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs):
        text = clean(paragraph.text)
        if not text:
            continue
        paragraphs.append(
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else "",
                "text": text,
            }
        )

    tables = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            rows.append([clean(cell.text) for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})

    payload = {
        "source": source.name,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "section_count": len(document.sections),
        "paragraphs": paragraphs,
        "tables": tables,
    }

    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    styles = {}
    for paragraph in paragraphs:
        styles[paragraph["style"]] = styles.get(paragraph["style"], 0) + 1

    print(
        json.dumps(
            {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "section_count": len(document.sections),
                "styles": styles,
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
